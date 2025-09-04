from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()
# doc.add_heading('示例表格', level=1)
# 添加标题段落
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中
run = title.add_run('示例表格')
run.bold = True                               # 加粗
run.font.size = Pt(16)                        # 字号，可根据需要调整
run.font.name = '黑体'                        # 设置西文字体（部分系统有效）
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), '黑体')   # 设置中文字体为黑体

# 在标题下添加普通文本行
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中
sub_run = subtitle.add_run('数据内容清单')
sub_run.bold = False                             # 不加粗
sub_run.font.size = Pt(12)                       # 可以调小字号
sub_run.font.name = '宋体'                        # 西文字体
sub_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 中文字体为宋体

# 添加“申请单位”行
applicant = doc.add_paragraph()
applicant.alignment = WD_ALIGN_PARAGRAPH.LEFT   # 左对齐
app_run = applicant.add_run('申请单位')
app_run.bold = False
app_run.font.size = Pt(12)
app_run.font.name = '宋体'
app_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 创建表格，初始一行表头
table = doc.add_table(rows=1, cols=4)
headers = ['序列号', '省份', '城市', '图层']

# 设置表头加粗居中
# hdr_cells = table.rows[0].cells
# for i, cell in enumerate(hdr_cells):
#     paragraph = cell.paragraphs[0]
#     paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
#     run = paragraph.add_run(headers[i])
#     run.bold = True
#     run.font.size = Pt(11)

# 示例数据，每行字典，城市可能多于1个
data = [
    {'序列号': '1', '省份': '广东省', '城市': ['广州', '深圳'], '图层': 'Layer1'},
    {'序列号': '2', '省份': '浙江省', '城市': ['杭州', '宁波', '温州'], '图层': 'Layer2'},
]

for row in data:
    cities = row['城市']
    n = len(cities)

    # 先插入所有城市行
    new_rows = [table.add_row() for _ in range(n)]
    for i, r in enumerate(new_rows):
        r.cells[0].text = str(row['序列号'])
        r.cells[1].text = row['省份']
        r.cells[2].text = cities[i]
        r.cells[3].text = row['图层']

    # 设置内容居中
    for r in new_rows:
        for cell in r.cells:
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(11)

    # 合并第一列、第二列、第四列
    # first_col_cells = [r.cells[0] for r in new_rows]
    # second_col_cells = [r.cells[1] for r in new_rows]
    # fourth_col_cells = [r.cells[3] for r in new_rows]

    # first_col_cells[0].merge(first_col_cells[-1])
    # second_col_cells[0].merge(second_col_cells[-1])
    # fourth_col_cells[0].merge(fourth_col_cells[-1])

      # 合并第1、2、4列，并保留内容在第一行
    for col_idx in [0, 1, 3]:
        cells_to_merge = [r.cells[col_idx] for r in new_rows]
        first_cell = cells_to_merge[0]
        # 保存内容
        content = first_cell.text
        merged_cell = first_cell.merge(cells_to_merge[-1])
        merged_cell.text = content
        # 居中
        merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
# 保存文档
doc.save('merged_table_after_insert.docx')
