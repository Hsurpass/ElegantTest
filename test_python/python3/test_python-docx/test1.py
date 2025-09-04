from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_VERTICAL, WD_ALIGN_PARAGRAPH

# 示例数据
data = [
    ['安徽省', '安庆市', '区码1', '100KM'],
    ['安徽省', '安庆市', '区码2', '100KM'],
    ['安徽省', '亳州市', '区码3', '200KM'],
    ['安徽省', '亳州市', '区码4', '200KM'],
    ['北京市', '北京市', '区码5', '300KM'],
]

headers = ['省份', '城市', '区码', '里程']

doc = Document()
table = doc.add_table(rows=1, cols=len(headers))
table.style = 'Table Grid'

# 设置表头
for idx, header in enumerate(headers):
    cell = table.cell(0, idx)
    cell.text = header
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.runs[0]
        run.font.bold = True
        run.font.size = Pt(12)

# 插入数据行
for row in data:
    cells = table.add_row().cells
    for idx, val in enumerate(row):
        cells[idx].text = str(val)
        cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in cells[idx].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.runs[0]
            run.font.size = Pt(11)

def merge_same_cells(table, col_idx):
    start = 1
    for i in range(2, len(table.rows)):
        if table.cell(i, col_idx).text != table.cell(start, col_idx).text:
            if i - start > 1:
                table.cell(start, col_idx).merge(table.cell(i-1, col_idx))
            start = i
    # 合并最后一组
    if len(table.rows) - start > 1:
        table.cell(start, col_idx).merge(table.cell(len(table.rows)-1, col_idx))

# 合并第1、2、4列（索引0,1,3）
for col in [0, 1, 3]:
    merge_same_cells(table, col)

doc.save('word表格合并示例.docx')